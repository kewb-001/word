# coding:utf-8

import os
import win32com
from win32com.client import Dispatch, constants
from docx import Document
import openpyxl


def parse_docx(f):
    """
    读取docx，匹配对应的内容
    """
    d = Document(f)
    t = d.tables[-1]
    # t_next = d.tables[-1]
    # title= t.cell(0,1).text//比如说标题在表格中位于（0，1）
    name = t.cell(0, 1).text  # 1
    Gender = t.cell(1, 3).text  # 1
    Race = t.cell(1, 5).text  # 1
    School = t.cell(2, 1).text  # 1
    Edu_back = t.cell(2, 4).text  # 1
    Major = t.cell(3, 1).text  # 1
    Grade = t.cell(3, 4).text  # 1
    Poli_Status = t.cell(4, 1).text  # 1
    Place = t.cell(4, 3).text  # 15
    ID_card = t.cell(5, 1).text  # 1
    Date_birth = t.cell(5, 4).text  # 1
    Telephone = t.cell(6, 1).text  # 1
    E_mail = t.cell(6, 4).text  # 1
    QQ = t.cell(7, 1).text  # 1
    Address = t.cell(7, 4).text  # 1
    #  Other=t_next.cell(1,1).text
    print(name, Gender, Race, School, Edu_back, Major, Grade, Poli_Status, Place, Date_birth, Telephone, E_mail, QQ,
          Address)


def write_excel(file):
    pass


''' 上述函数主要实现文件的读取 '''
if __name__ == "__main__":
    # w = win32com.client.Dispatch('Word.Application')
    # 遍历文件
    PATH = r"C:\Users\v_vwbiaoke\Desktop\123.docx"  # windows文件路径
    parse_docx(PATH)
