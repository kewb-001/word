import openpyxl
from docx import Document  # 导入读取word模块
import time

time = time.strftime("%Y%m", time.localtime())
time1 = time


# 打开word文件,读取关键信息
def Handling_Information():
    import time
    # dlg = win32ui.CreateFileDialog(1)  # 1代表Ture，表示打开文件对话框
    # dlg.SetOFNInitialDir(r'C:\Users\14325\PycharmProjects\word\input')  # 设置打开文件对话框中的初始显示目录
    # dlg.DoModal()  # 等待用户选择文件
    # WordName = dlg.GetPathName()  # 获取选择的文件名称
    # WordPath = os.path.join(os.getcwd(), WordName)  # 获取选择的文件路径
    # print(WordPath)
    path_word = 'input/GST-Form-SD-002 通用测试申请表.docx'
    # doc = Document(path_word)  # 读取word文档
    doc = Document(path_word)

    # doc = Document(WordPath)
    tables = doc.tables  # 读取word中的表格
    table = tables[0]
    header = doc.sections[0].header.paragraphs[0].text  # 读取表头，识别哪种提取规则
    if header == 'General Testing Requisition Form':
        #  提取表头的指定表格值
        header_form_Report_No = doc.sections[0].header.tables[0].cell(1, 1).text  # Report No.报告号
        header_form_Start_Date = doc.sections[0].header.tables[0].cell(1, 3).text  # Start Date开案日期
        header_form_Sales_Rep = doc.sections[0].header.tables[0].cell(2, 3).text  # Sales Rep.销售工程师
        header_form_Due_Date = doc.sections[0].header.tables[0].cell(1, 5).text  # Due Date完成日期
        header_form_CS_No = doc.sections[0].header.tables[0].cell(2, 5).text  # CS No.客服编号

        #  提取正文中的指定表格值
        Applicant_Name_English = table.cell(2, 4).text  # 申请公司英文名
        Applicant_Name_Chinese = table.cell(3, 4).text  # 申请公司中文名
        Sample_name = table.cell(11, 0).text  # 样品名称
        Test_Items = table.cell(11, 7).text  # 测试项目
        #  写入Excel
        path_excel = 'input/123.xlsx'
        wb = openpyxl.load_workbook(path_excel)
        sheetname = wb[time1]
        nrows = sheetname.max_row
        sheetname.cell(nrows + 1, 1, Applicant_Name_English)  # 客户名称--cell(行，列，值）
        sheetname.cell(nrows + 1, 2, header_form_Start_Date)  # 开案时间
        sheetname.cell(nrows + 1, 3, header_form_Due_Date)  # 合同完成日期
        sheetname.cell(nrows + 1, 4, )  # 合同完成具体时间
        sheetname.cell(nrows + 1, 5, header_form_Sales_Rep)  # 客户归属
        sheetname.cell(nrows + 1, 6, header_form_CS_No)  # 客服
        sheetname.cell(nrows + 1, 7, )  # 报价单号
        sheetname.cell(nrows + 1, 8, header_form_Report_No)  # 报告号
        sheetname.cell(nrows + 1, 9, Sample_name)  # 样品名称
        sheetname.cell(nrows + 1, 10, Test_Items)  # 测试项目
        wb.save(path_excel)
    elif header == 'General Testing Requisition For':
        print('改')
    else:
        print('sda')


if __name__ == '__main__':
    Handling_Information()
