# -*- coding: UTF-8 -*-
from docx import Document


# 读取docx文档内容并转存为txt
def readSpecTable(filename, specText):
    document = Document(filename)
    paragraphs = document.paragraphs
    allTables = document.tables
    specText = specText.encode('utf-8').decode('utf-8')
    for aPara in paragraphs:
        if aPara.text == specText:
            ele = aPara._p.getnext()
            while (ele.tag != '' and ele.tag[-3:] != 'tbl'):
                ele = ele.getnext()
            if ele.tag != '':
                for aTable in allTables:
                    if aTable._tbl == ele:
                        for i in range(len(aTable.rows)):
                            for j in range(len(aTable.columns)):
                                # print(aTable.cell(i, j).text)
                                f = open(r"C:\Users\v_vwbiaoke\Desktop\123.txt", 'a')
                                f.write(aTable.cell(i, j).text)


def Extract_Contents(txt_filename):
    pass


if __name__ == '__main__':
    readSpecTable(r"C:\Users\v_vwbiaoke\Desktop\123.docx", '正式党员信息登记表')
