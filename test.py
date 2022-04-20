import os

import docx
from win32com import client as wc


def find_doc(raw_path):
    """
    : 获取当前文件夹下所有.doc文件路径，包含子文件夹下
    : param raw_path: 传入待处理的文件夹
    : return: doc_file_list，绝对路径
    """
    doc_file_list = []
    for home, dirs, files in os.walk(raw_path):
        for filename in files:
            if filename.endswith('.doc'):
                doc_file_list.append(os.path.join(home, filename))
    return doc_file_list


def save_doc_to_docx(root_path):
    """
    : param raw_path: 传入和传出文件夹的路径
    : return: None
    """
    word = wc.Dispatch("Word.Application")
    # 不能使用相对路径，使用绝对路径
    doc_files = find_doc(root_path)
    for doc_file in doc_files:
        # ~$是为了排除临时文件
        if not doc_file.startswith('~$'):
            # 打开文件
            # print(doc_file)
            doc = word.Documents.Open(doc_file)
            # 分割文件名
            rename = os.path.splitext(doc_file)
            # 将文件另存为docx
            doc.SaveAs(rename[0] + '.docx', 12)  # 12表示docx格式
            doc.Close()
    word.Quit()


def read_docx():
    path = r"C:\Users\v_vwbiaoke\Desktop\123.docx"  # 文件路径
    document = docx.Document(path)  # 读入文件
    for para in document.paragraphs:  # 读取非表格内容
        print(para.text)
    tables = document.tables  # 获取文件中的表格集
    table = tables[0]  # 获取文件中的第一个表格
    for i in range(0, len(table.rows)):  # 从表格第二行开始循环读取表格数据
        result = table.cell(i, 0).text + "" + table.cell(i, 1).text + table.cell(i, 2).text + table.cell(i, 3).text
        # cell(i,0)表示第(i+1)行第1列数据，以此类推
        print(result)
        f = open(r"C:\Users\v_vwbiaoke\Desktop\123.txt", 'a')
        f.write(result)


if __name__ == '__main__':
    # find_doc(r'C:\Users\v_vwbiaoke\Desktop')
    # save_doc2docx(r'C:\Users\v_vwbiaoke\Desktop')
    read_docx()
