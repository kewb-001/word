import re
import os
from docx import Document

filepath = r'C:\Users\v_vwbiaoke\Desktop'  # 文件夹路径
file_list = os.listdir(filepath)  # 获取文件夹下所有文件的名称，包括后缀


# filepath1 = os.path.join(filepath,file_list[0])
# file_list

# document = Document(filepath1)


# 采集docx文件信息
def docxInfo(addr):
    document = Document(addr)

    info = {'name': [],
            'ID': [],
            'adress': [],
            'role': [],
            'case_date': [],
            'boat': [],
            'index_doc': [],
            'fpath': []}  # 增加字典空間

    lines = [0 for i in range(len(document.paragraphs))]
    k = 0
    name = []
    ID = []
    adress = []
    role = []
    case_date = []
    boat = []
    index_doc = []
    fpath = []  # 初始化定义
    # 使用正则时候，将你需要提取信息前后的中文转为unicode
    for paragraph in document.paragraphs:
        lines[k] = paragraph.text
        w1 = "\u59d3\u540d\u000d\u000a"  # "姓名"的unicode
        w2 = '\uFF0C'  # ","的unicode
        w3 = '\u8EAB\u4EFD\u8BC1'  # "身份证"的unicode
        w4 = '\u6237\u7C4D'  # "户籍的unicode
        w5 = '\u6587\u5316\uFF0C'  # "文化，"unicode
        w6 = '\u3002'  # "。"的unicode
        w7 = '\u56E0'  # "因"的unicode
        w8 = '\u201C'  # "“"的unicode
        w9 = '\u201D'  # "”"的unicode
        pat_name = re.compile(w1 + '(.*?)' + w2, re.S)
        result_name = pat_name.findall(lines[k])
        pat_ID = re.compile(w3 + '(.*?)' + w6, re.S)
        result_ID = pat_ID.findall(lines[k])
        pat_adress = re.compile(w4 + '(.*?)' + w2, re.S)
        result_adress = pat_adress.findall(lines[k])
        pat_role = re.compile(w5 + '(.*?)' + w2, re.S)
        result_role = pat_role.findall(lines[k])
        pat_case_date = re.compile(w6 + '(.*?)' + w7, re.S)
        result_case_date = pat_case_date.findall(lines[k])
        pat_boat = re.compile(w8 + '(.*?)' + w9, re.S)
        result_boat = pat_boat.findall(lines[k])
        name.append(result_name)
        ID.append(result_ID)
        adress.append(result_adress)
        role.append(result_role)
        case_date.append(result_case_date)
        boat.append(result_boat)
        index_doc.append(k)
        fpath.append(addr)
        info['name'] = name
        info['ID'] = ID
        info['adress'] = adress
        info['role'] = role
        info['case_date'] = case_date
        info['boat'] = boat
        info['fpath'] = fpath
        k = k + 1
    info['index_doc'] = index_doc

    return info


# if __name__ == '__main__':
#	print(docxInfo(filepath1))


import pandas as pd


# 将采集的信息（dict格式）转换为表格保存
def export_excel(export):
    # 将字典转换为DataFrame
    pf = pd.DataFrame.from_dict(export)
    # 指定字段顺序
    order = ['name', 'ID', 'adress', 'role', 'case_date', 'boat', 'index_doc', 'fpath']
    pf = pf[order]
    # 将列名替换为中文
    columns_map = {
        'name': '姓名',
        'ID': '身份证号',
        'adress': '户籍所在地',
        'role': '角色',
        'case_date': '日期',
        'boat': '船只',
        'index_doc': '截段索引',
        'fpath': '路径'
    }
    pf.rename(columns=columns_map, inplace=True)
    return pf


# 判断文件后缀
def endWith(s, *endstring):
    array = map(s.endswith, endstring)
    if True in array:
        return True
    else:
        return False


if __name__ == '__main__':
    i = 0
    result_df = pd.DataFrame()
    for file in file_list:
        file_path = os.path.join(filepath, file_list[i])
        i += 1
        if endWith(file, '.docx'):
            file_info1 = docxInfo(file_path)
            # 将分析完成的列表导出为pf
            xeport_data_1 = export_excel(file_info1)  # 一个pf
            # 删除/选取某列含有特定数值'[]'的行（還沒調試成功）
            # xeport_data_1=xeport_data_1[xeport_data_1['身份证号'].isin([[]])]
            result_df = pd.concat([result_df, xeport_data_1], axis=0, sort=False)
    filename = file.split(".")
    result_df.to_excel(r'C:\Users\v_vwbiaoke\Desktop\{}.xlsx'.format(filename[0]), encoding='utf-8', index=False)
